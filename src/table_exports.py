from __future__ import annotations

import json
import re
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from openpyxl import Workbook, load_workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter


def _typographic_negative_signs(text: str) -> str:
    return re.sub(r"(?<![A-Za-z])-(?=\d)", "−", str(text))


def _write_cell(cell, text: str, *, bold: bool = False, centred: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    for run in list(paragraph.runs):
        paragraph._p.remove(run._element)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if centred else WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1
    run = paragraph.add_run(_typographic_negative_signs(text))
    _set_run_font(run, "Arial", 9, bold=bold)


def _add_physician_clustering_body(document: Document, data: pd.DataFrame) -> None:
    patient = data[data["Section"] == "A. Patient-reported outcomes"]
    icc = data[data["Section"] == "B. Physician-level ICCs"]
    row_count = 4 + len(patient) + len(icc)
    table = document.add_table(rows=row_count, cols=5)
    table.style = "Normal Table"
    _set_table_geometry(table, [2500, 1650, 1500, 2300, 1300])

    section_b_index = 2 + len(patient)
    header_b_index = section_b_index + 1
    table.rows[0].cells[0].merge(table.rows[0].cells[4])
    table.rows[section_b_index].cells[0].merge(table.rows[section_b_index].cells[4])
    table.rows[header_b_index].cells[1].merge(table.rows[header_b_index].cells[4])
    for row_index in range(header_b_index + 1, row_count):
        table.rows[row_index].cells[1].merge(table.rows[row_index].cells[4])

    _write_cell(table.rows[0].cells[0], "A. Patient-reported outcomes", bold=True)
    for cell, label in zip(table.rows[1].cells, [
        "Outcome",
        "Conventional patient-level estimate",
        "Physician-clustered estimate",
        "Physician-clustered 95% CI",
        "Physician-clustered P value",
    ]):
        _write_cell(cell, label, bold=True, centred=True)

    for row_index, values in enumerate(patient.itertuples(index=False), start=2):
        row_values = [values.Outcome, values._2, values._3, values._4, values._5]
        for column, (cell, value) in enumerate(zip(table.rows[row_index].cells, row_values)):
            _write_cell(cell, value, centred=column > 0)

    _write_cell(
        table.rows[section_b_index].cells[0], "B. Physician-level ICCs", bold=True
    )
    _write_cell(
        table.rows[header_b_index].cells[0], "Outcome", bold=True, centred=True
    )
    _write_cell(
        table.rows[header_b_index].cells[1], "Physician-level ICC", bold=True, centred=True
    )
    for row_index, values in enumerate(
        icc.itertuples(index=False), start=header_b_index + 1
    ):
        _write_cell(table.rows[row_index].cells[0], values.Outcome)
        _write_cell(table.rows[row_index].cells[1], values._6, centred=True)

    for row_index, row in enumerate(table.rows):
        cells = []
        seen = set()
        for cell in row.cells:
            marker = id(cell._tc)
            if marker not in seen:
                seen.add(marker)
                cells.append(cell)
        for cell in cells:
            top = (
                ("12", "000000") if row_index == 0
                else (("8", "000000") if row_index == section_b_index else None)
            )
            bottom = (
                ("12", "000000") if row_index == row_count - 1
                else (("8", "000000") if row_index in {1, header_b_index} else None)
            )
            _set_cell_border(cell, top=top, bottom=bottom)


def write_physician_clustering_word(
    table_dir: Path,
    output_dir: Path,
) -> Path:
    output_dir.mkdir(parents=True, exist_ok=True)
    data = pd.read_csv(
        table_dir / "rct_physician_clustering_outcomes.csv",
        dtype=str,
        keep_default_na=False,
    )
    document = Document()
    section = document.sections[0]
    section.page_width = Inches(8.2680556)
    section.page_height = Inches(11.6930556)
    section.left_margin = Inches(0.8861111)
    section.right_margin = Inches(0.8861111)
    section.top_margin = Inches(0.8270833)
    section.bottom_margin = Inches(0.7875)
    _add_physician_clustering_body(document, data)
    path = output_dir / PHYSICIAN_ICC_WORD_DOCUMENT
    document.save(path)
    return path


EXCEL_WORKBOOKS = {
    "rct_results.xlsx": [
        ("Baseline characteristics", "rct_baseline_characteristics.csv"),
        ("RCT outcomes", "rct_outcomes.csv"),
        ("GCA safety and implementation", "rct_ai_safety_and_implementation.csv"),
    ],
    "multicentre_characteristics.xlsx": [
        ("Multicentre characteristics", "multicentre_characteristics.csv"),
    ],
    "multicentre_and_seniority_results.xlsx": [
        ("RCT seniority", "rct_seniority_effects.csv"),
        ("Overall paired effects", "multicentre_overall_paired_effects.csv"),
        ("Centre-specific effects", "multicentre_centre_specific_effects.csv"),
        ("Paired-effect sensitivity", "multicentre_paired_effect_sensitivity.csv"),
        ("Centre heterogeneity", "multicentre_centre_heterogeneity.csv"),
        ("PDQI-9 domains", "multicentre_pdqi_domains.csv"),
        ("Multicentre seniority", "multicentre_seniority_effects.csv"),
        ("PDQI-9 seniority gaps", "multicentre_pdqi_seniority_gaps.csv"),
        ("Draft safety", "multicentre_draft_safety.csv"),
        ("Safety inference", "multicentre_safety_inference.csv"),
    ],
    "Sensitivity_Analyses.xlsx": [
        ("RCT ordinal", "rct_ordinal_sensitivity.csv"),
        ("Multicentre ordinal", "multicentre_ordinal_gee_sensitivity.csv"),
    ],
    "physician_clustering_outcomes.xlsx": [
        ("Physician clustering", "rct_physician_clustering_outcomes.csv"),
    ],
}


PHYSICIAN_CLUSTER_WORD_DOCUMENT = "physician_clustered_sensitivity.docx"
PHYSICIAN_ICC_WORD_DOCUMENT = "physician_clustering_outcomes.docx"


def _excel_table_typography(csv_path: Path) -> tuple[str, float, float]:
    name = csv_path.name
    if name in {
        "rct_baseline_characteristics.csv",
        "rct_outcomes.csv",
        "rct_ai_safety_and_implementation.csv",
    }:
        return "Times New Roman", 7.5, 8.0 if name == "rct_ai_safety_and_implementation.csv" else 7.5
    if name.startswith("multicentre_") and name != "multicentre_characteristics.csv":
        return "Arial", 7.5, 7.5
    if name == "rct_seniority_effects.csv":
        return "Arial", 7.5, 7.5
    return "Arial", 9.0, 9.0


def _add_excel_sheet(workbook: Workbook, sheet_name: str, csv_path: Path) -> None:
    data = pd.read_csv(csv_path, dtype=str, keep_default_na=False)
    sheet = workbook.create_sheet(sheet_name)
    sheet.sheet_view.showGridLines = False
    sheet.sheet_properties.pageSetUpPr.fitToPage = True
    sheet.page_setup.orientation = "landscape"
    sheet.page_setup.fitToWidth = 1
    sheet.page_setup.fitToHeight = 0
    sheet.freeze_panes = "A2"

    columns = list(data.columns)
    sheet.append(columns)
    for row in data.itertuples(index=False, name=None):
        sheet.append([str(value) for value in row])

    black = "000000"
    font_name, body_size, header_size = _excel_table_typography(csv_path)
    section_fill = PatternFill("solid", fgColor="F3F4F6")
    medium = Side(style="medium", color=black)
    thin = Side(style="thin", color=black)

    for cell in sheet[1]:
        cell.font = Font(name=font_name, size=header_size, bold=True)
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        cell.border = Border(top=medium, bottom=thin)
    sheet.row_dimensions[1].height = 30

    for row_index in range(2, sheet.max_row + 1):
        values = [
            str(sheet.cell(row_index, column).value or "")
            for column in range(1, sheet.max_column + 1)
        ]
        section_row = bool(values[0]) and all(not value for value in values[1:])
        for column_index in range(1, sheet.max_column + 1):
            cell = sheet.cell(row_index, column_index)
            cell.font = Font(
                name=font_name,
                size=header_size if section_row else body_size,
                bold=section_row,
            )
            if section_row:
                cell.fill = section_fill
            cell.alignment = Alignment(
                horizontal="left" if column_index == 1 else "center",
                vertical="top",
                wrap_text=True,
            )
            cell.border = Border(bottom=medium if row_index == sheet.max_row else Side(style=None))

    for column_index, column in enumerate(columns, start=1):
        values = [str(column)] + [str(value) for value in data.iloc[:, column_index - 1].tolist()]
        width = max(len(value) for value in values) + 2
        if column_index == 1:
            width = min(max(width, 24), 48)
        else:
            width = min(max(width, 12), 36)
        sheet.column_dimensions[get_column_letter(column_index)].width = width

    sheet.auto_filter.ref = f"A1:{get_column_letter(sheet.max_column)}{sheet.max_row}"
    sheet.print_title_rows = "1:1"
    sheet.print_area = f"A1:{get_column_letter(sheet.max_column)}{sheet.max_row}"


def write_excel_workbooks(table_dir: Path, output_dir: Path) -> list[Path]:
    output_dir.mkdir(parents=True, exist_ok=True)
    generated = []
    for filename, sheets in EXCEL_WORKBOOKS.items():
        workbook = Workbook()
        workbook.remove(workbook.active)
        for sheet_name, csv_name in sheets:
            _add_excel_sheet(workbook, sheet_name, table_dir / csv_name)
        path = output_dir / filename
        workbook.save(path)
        generated.append(path)
    return generated


def _set_run_font(run, name: str, size: float, bold: bool = False) -> None:
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)


def _set_cell_border(cell, top: tuple[str, str] | None = None, bottom: tuple[str, str] | None = None) -> None:
    properties = cell._tc.get_or_add_tcPr()
    borders = properties.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        properties.append(borders)
    specifications = {
        "top": top,
        "bottom": bottom,
        "left": None,
        "right": None,
    }
    for edge, specification in specifications.items():
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        if specification is None:
            element.set(qn("w:val"), "nil")
        else:
            size, color = specification
            element.set(qn("w:val"), "single")
            element.set(qn("w:sz"), size)
            element.set(qn("w:space"), "0")
            element.set(qn("w:color"), color)


def _set_table_geometry(table, widths: list[int]) -> None:
    table.autofit = False
    properties = table._tbl.tblPr
    width = properties.find(qn("w:tblW"))
    if width is None:
        width = OxmlElement("w:tblW")
        properties.append(width)
    width.set(qn("w:type"), "dxa")
    width.set(qn("w:w"), str(sum(widths)))
    layout = properties.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        properties.append(layout)
    layout.set(qn("w:type"), "fixed")
    indent = properties.find(qn("w:tblInd"))
    if indent is None:
        indent = OxmlElement("w:tblInd")
        properties.append(indent)
    indent.set(qn("w:type"), "dxa")
    indent.set(qn("w:w"), "0")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for value in widths:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(value))
        grid.append(column)
    for row in table.rows:
        cant_split = OxmlElement("w:cantSplit")
        row._tr.get_or_add_trPr().append(cant_split)
        for cell, value in zip(row.cells, widths):
            cell.width = Inches(value / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell_properties = cell._tc.get_or_add_tcPr()
            cell_width = cell_properties.find(qn("w:tcW"))
            if cell_width is None:
                cell_width = OxmlElement("w:tcW")
                cell_properties.append(cell_width)
            cell_width.set(qn("w:type"), "dxa")
            cell_width.set(qn("w:w"), str(value))


def _add_three_line_table(document: Document, data: pd.DataFrame) -> None:
    table = document.add_table(rows=1, cols=len(data.columns))
    table.style = "Normal Table"
    header = table.rows[0]
    header_properties = header._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    header_properties.append(repeat)
    for column, label in enumerate(data.columns):
        cell = header.cells[column]
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(str(label))
        _set_run_font(run, "Arial", 9, bold=True)
        _set_cell_border(cell, top=("12", "000000"), bottom=("8", "000000"))

    for values in data.itertuples(index=False, name=None):
        row = table.add_row()
        for column, value in enumerate(values):
            cell = row.cells[column]
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT if column < 2 else WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1
            text = str(value)
            if str(data.columns[column]).startswith("Effect estimate"):
                text = _typographic_negative_signs(text)
            run = paragraph.add_run(text)
            _set_run_font(run, "Arial", 9)
            _set_cell_border(cell)
    for cell in table.rows[-1].cells:
        _set_cell_border(cell, bottom=("12", "000000"))
    _set_table_geometry(table, [2700, 3000, 3550])


def write_physician_cluster_word_tables(table_dir: Path, output_dir: Path) -> Path:
    output_dir.mkdir(parents=True, exist_ok=True)
    rct = pd.read_csv(
        table_dir / "rct_physician_clustered_sensitivity.csv",
        dtype=str,
        keep_default_na=False,
    )
    external = pd.read_csv(
        table_dir / "multicentre_physician_clustered_sensitivity.csv",
        dtype=str,
        keep_default_na=False,
    )
    document = Document()
    section = document.sections[0]
    section.page_width = Inches(8.2680556)
    section.page_height = Inches(11.6930556)
    section.left_margin = Inches(0.8861111)
    section.right_margin = Inches(0.8861111)
    section.top_margin = Inches(0.8270833)
    section.bottom_margin = Inches(0.7875)

    _add_three_line_table(document, rct)
    document.add_page_break()
    _add_three_line_table(document, external)
    path = output_dir / PHYSICIAN_CLUSTER_WORD_DOCUMENT
    document.save(path)
    return path


def _cell_border_value(cell, edge: str) -> str | None:
    borders = cell._tc.get_or_add_tcPr().find(qn("w:tcBorders"))
    if borders is None:
        return None
    element = borders.find(qn(f"w:{edge}"))
    return None if element is None else element.get(qn("w:val"))


def _validate_physician_cluster_word(path: Path, table_dir: Path) -> tuple[bool, bool]:
    document = Document(path)
    sources = [
        "rct_physician_clustered_sensitivity.csv",
        "multicentre_physician_clustered_sensitivity.csv",
    ]
    if len(document.tables) != len(sources):
        return False, False
    content_matches = True
    formatting_matches = True
    for table, source in zip(document.tables, sources):
        expected = pd.read_csv(table_dir / source, dtype=str, keep_default_na=False)
        expected_rows = [list(expected.columns)] + [
            [
                _typographic_negative_signs(value)
                if str(expected.columns[column]).startswith("Effect estimate")
                else str(value)
                for column, value in enumerate(row)
            ]
            for row in expected.itertuples(index=False, name=None)
        ]
        observed_rows = [[cell.text for cell in row.cells] for row in table.rows]
        content_matches = content_matches and observed_rows == expected_rows
        grid = [int(column.get(qn("w:w"))) for column in table._tbl.tblGrid.gridCol_lst]
        formatting_matches = formatting_matches and grid == [2700, 3000, 3550]
        for row_index, row in enumerate(table.rows):
            for cell in row.cells:
                left = _cell_border_value(cell, "left")
                right = _cell_border_value(cell, "right")
                top = _cell_border_value(cell, "top")
                bottom = _cell_border_value(cell, "bottom")
                formatting_matches = formatting_matches and left in {None, "nil"} and right in {None, "nil"}
                if row_index == 0:
                    formatting_matches = formatting_matches and top == "single" and bottom == "single"
                elif row_index == len(table.rows) - 1:
                    formatting_matches = formatting_matches and top in {None, "nil"} and bottom == "single"
                else:
                    formatting_matches = formatting_matches and top in {None, "nil"} and bottom in {None, "nil"}
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        formatting_matches = formatting_matches and run.font.name == "Arial"
                        formatting_matches = formatting_matches and run.font.size is not None and abs(run.font.size.pt - 9) < 1e-9
    content_matches = content_matches and all(
        not paragraph.text.strip() for paragraph in document.paragraphs
    )
    return content_matches, formatting_matches


def _validate_physician_clustering_word(path: Path, table_dir: Path) -> tuple[bool, bool]:
    document = Document(path)
    if len(document.tables) != 1:
        return False, False
    data = pd.read_csv(
        table_dir / "rct_physician_clustering_outcomes.csv",
        dtype=str,
        keep_default_na=False,
    )
    patient = data[data["Section"] == "A. Patient-reported outcomes"]
    icc = data[data["Section"] == "B. Physician-level ICCs"]
    table = document.tables[0]
    expected_rows = 4 + len(patient) + len(icc)
    content_matches = len(table.rows) == expected_rows
    if not content_matches:
        return False, False

    content_matches = (
        table.rows[0].cells[0].text == "A. Patient-reported outcomes"
        and table.rows[2 + len(patient)].cells[0].text == "B. Physician-level ICCs"
    )
    for row_index, values in enumerate(patient.itertuples(index=False), start=2):
        expected = [
            values.Outcome, values._2, values._3, values._4, values._5,
        ]
        observed = [cell.text for cell in table.rows[row_index].cells[:5]]
        content_matches = content_matches and observed == [
            _typographic_negative_signs(value) for value in expected
        ]
    start = 4 + len(patient)
    for row_index, values in enumerate(icc.itertuples(index=False), start=start):
        content_matches = (
            content_matches
            and table.rows[row_index].cells[0].text == values.Outcome
            and table.rows[row_index].cells[1].text == values._6
        )

    grid = [int(column.get(qn("w:w"))) for column in table._tbl.tblGrid.gridCol_lst]
    formatting_matches = grid == [2500, 1650, 1500, 2300, 1300]
    for row in table.rows:
        seen = set()
        for cell in row.cells:
            marker = id(cell._tc)
            if marker in seen:
                continue
            seen.add(marker)
            formatting_matches = (
                formatting_matches
                and _cell_border_value(cell, "left") in {None, "nil"}
                and _cell_border_value(cell, "right") in {None, "nil"}
            )
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    formatting_matches = (
                        formatting_matches
                        and run.font.name == "Arial"
                        and run.font.size is not None
                        and abs(run.font.size.pt - 9) < 1e-9
                    )
    content_matches = content_matches and all(
        not paragraph.text.strip() for paragraph in document.paragraphs
    )
    return content_matches, formatting_matches


def validate_exports(excel_dir: Path, word_dir: Path | None, qa_dir: Path) -> bool:
    rows = []
    for filename, sheets in EXCEL_WORKBOOKS.items():
        path = excel_dir / filename
        exists = path.exists() and path.stat().st_size > 0
        observed_sheets = []
        nonempty = False
        content_matches = False
        formatting_matches = False
        if exists:
            workbook = load_workbook(path, read_only=False, data_only=False)
            observed_sheets = workbook.sheetnames
            nonempty = all(workbook[name].max_row >= 2 and workbook[name].max_column >= 2 for name, _ in sheets)
            content_matches = True
            formatting_matches = True
            for sheet_name, csv_name in sheets:
                sheet = workbook[sheet_name]
                expected = pd.read_csv(excel_dir.parent / "tables" / csv_name, dtype=str, keep_default_na=False)
                observed_columns = [str(sheet.cell(1, column).value or "") for column in range(1, sheet.max_column + 1)]
                observed_rows = [
                    [str(sheet.cell(row, column).value or "") for column in range(1, sheet.max_column + 1)]
                    for row in range(2, sheet.max_row + 1)
                ]
                expected_rows = [[str(value) for value in row] for row in expected.itertuples(index=False, name=None)]
                content_matches = content_matches and observed_columns == list(expected.columns)
                content_matches = content_matches and observed_rows == expected_rows
                content_matches = content_matches and not sheet.merged_cells.ranges
                formatting_matches = formatting_matches and sheet.sheet_view.showGridLines is False
                formatting_matches = formatting_matches and str(sheet.freeze_panes) == "A2"
            workbook.close()
        expected_sheets = [name for name, _ in sheets]
        rows.append({
            "Format": "Excel",
            "File": filename,
            "Observed": "; ".join(observed_sheets),
            "Expected": "; ".join(expected_sheets),
            "Status": "PASS" if exists and observed_sheets == expected_sheets and nonempty and content_matches and formatting_matches else "FAIL",
        })

    if word_dir is not None:
        path = word_dir / PHYSICIAN_CLUSTER_WORD_DOCUMENT
        observed_tables = len(Document(path).tables) if path.exists() else 0
        content_matches, formatting_matches = (
            _validate_physician_cluster_word(path, excel_dir.parent / "tables")
            if path.exists() else (False, False)
        )
        rows.append({
            "Format": "Word",
            "File": PHYSICIAN_CLUSTER_WORD_DOCUMENT,
            "Observed": f"{observed_tables} tables; content={content_matches}; three_line_format={formatting_matches}",
            "Expected": "2 tables; content=True; three_line_format=True",
            "Status": "PASS" if observed_tables == 2 and content_matches and formatting_matches else "FAIL",
        })
        path = word_dir / PHYSICIAN_ICC_WORD_DOCUMENT
        observed_tables = len(Document(path).tables) if path.exists() else 0
        content_matches, formatting_matches = (
            _validate_physician_clustering_word(path, excel_dir.parent / "tables")
            if path.exists() else (False, False)
        )
        rows.append({
            "Format": "Word",
            "File": PHYSICIAN_ICC_WORD_DOCUMENT,
            "Observed": f"{observed_tables} table; content={content_matches}; three_line_format={formatting_matches}",
            "Expected": "1 table; content=True; three_line_format=True",
            "Status": "PASS" if observed_tables == 1 and content_matches and formatting_matches else "FAIL",
        })

    qa_dir.mkdir(parents=True, exist_ok=True)
    pd.DataFrame(rows).to_csv(qa_dir / "table_export_checks.csv", index=False, encoding="utf-8-sig")
    passed = all(row["Status"] == "PASS" for row in rows)
    manifest = {
        "excel_workbooks": sorted(path.name for path in excel_dir.glob("*.xlsx")),
        "word_documents": [] if word_dir is None else sorted(path.name for path in word_dir.glob("*.docx")),
        "passed": passed,
    }
    (qa_dir / "table_export_manifest.json").write_text(json.dumps(manifest, indent=2), encoding="utf-8")
    return passed
